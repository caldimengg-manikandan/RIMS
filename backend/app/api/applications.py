from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, BackgroundTasks, Form, Request
# Performance Serialization Fallback
try:
    import orjson
    from fastapi.responses import ORJSONResponse
except ImportError:
    from fastapi.responses import JSONResponse as ORJSONResponse
from sqlalchemy import or_, func, text, cast, String, extract, inspect as sa_inspect
from sqlalchemy.orm import Session, joinedload, selectinload, defer, load_only
from sqlalchemy.orm.exc import ObjectDeletedError
from app.core.storage import upload_file, get_signed_url, get_public_url
import os
import os
import json
import logging
import time
from datetime import datetime, timezone
from app.infrastructure.database import get_db, SessionLocal
from app.domain.models import User, Application, Job, ResumeExtraction, Interview, InterviewAnswer, ResumeExtractionVersion, InterviewReport
from app.domain.schemas import (
    ApplicationCreate,
    ApplicationStatusUpdate,
    ApplicationResponse,
    ApplicationDetailResponse,
    ApplicationSummaryResponse,
    ApplicationNotesUpdate,
    JobSummary,
    ResumeExtractionSummary,
    InterviewSummary,
    HasAppliedResponse,
    ApplicationListResponse,
)
from app.core.auth import get_current_user, get_current_hr, get_current_admin
from app.core.ownership import validate_hr_ownership
from app.services.ai_service import parse_resume_with_ai, extract_basic_candidate_info
from app.services.email_service import send_application_received_email, send_rejected_email, send_approved_for_interview_email
import secrets
from passlib.context import CryptContext
from datetime import datetime, timedelta
from sqlalchemy.exc import IntegrityError
import re
import hashlib

from typing import Optional, List, Union

logger = logging.getLogger(__name__)

# Persisted hint for HR (marker stripped in API responses). Heuristic also sets extraction_degraded.
RIMS_EXTRACTION_DEGRADED_MARKER = "[[rims:extraction_degraded]]"


def _strip_extraction_marker(notes: Optional[str]) -> Optional[str]:
    if not notes:
        return notes
    stripped = notes.replace(RIMS_EXTRACTION_DEGRADED_MARKER, "").strip()
    return stripped if stripped else None


def _append_extraction_degraded_marker(application: Application) -> None:
    current = application.hr_notes or ""
    if RIMS_EXTRACTION_DEGRADED_MARKER in current:
        return
    application.hr_notes = (current.rstrip() + "\n" + RIMS_EXTRACTION_DEGRADED_MARKER).strip()


def _heuristic_extraction_degraded(application: Application) -> bool:
    re = application.resume_extraction
    if not re:
        return False
    summary = re.summary or ""
    if "AI was unable to generate a summary for this resume." in summary:
        return True
    try:
        skills = json.loads(re.extracted_skills or "[]")
    except Exception:
        return False
    if skills == ["General Profile"]:
        return True
    # Avoid triggering a lazy-load of extracted_text (it can be large).
    # If extracted_text isn't loaded in the current query, we can't reliably apply this heuristic.
    try:
        insp = sa_inspect(re)
        if 'extracted_text' not in getattr(insp, "unloaded", set()):
            et = (re.extracted_text or "").strip()
            if et in ("Error extracting text.", "No readable text found."):
                return True
    except Exception:
        # Best-effort: if introspection fails, keep previous behavior as a fallback.
        et = (re.extracted_text or "").strip()
        if et in ("Error extracting text.", "No readable text found."):
            return True
    return False


def build_application_summary_response(application: Application, current_user_id: Optional[int] = None) -> ApplicationSummaryResponse:
    """Ultra-high performance summary builder using model_construct and manual mapping."""
    # 1. Manual mapping for nested models to ensure ZERO validation cost
    job_summary = None
    if application.job:
        job_summary = JobSummary.model_construct(
            id=application.job.id,
            job_id=application.job.job_id,
            title=application.job.title
        )

    re_summary = None
    if application.resume_extraction:
        re_summary = ResumeExtractionSummary.model_construct(
            id=application.resume_extraction.id,
            resume_score=application.resume_extraction.resume_score or 0.0,
            skill_match_percentage=application.resume_extraction.skill_match_percentage or 0.0,
            experience_level=application.resume_extraction.experience_level,
            summary=application.resume_extraction.summary,
            extracted_skills=application.resume_extraction.extracted_skills
        )

    int_summary = None
    if application.interview:
        int_summary = InterviewSummary.model_construct(
            id=application.interview.id,
            status=application.interview.status,
            overall_score=application.interview.overall_score or 0.0
        )

    # 2. Ownership Calculation (Simplified per Senior Backend Engineer request)
    is_owner = (application.hr_id == current_user_id) if current_user_id else False

    # Generate full photo URL from path — use public URL (no network call) for list view.
    # Signed URLs (which require a Supabase round-trip) are only needed on the detail page.
    photo_url = None
    if application.candidate_photo_path:
        photo_url = get_public_url(settings.supabase_bucket_id_photos, application.candidate_photo_path)

    # 3. Final Construction (model_construct bypasses all Pydantic validators)
    return ApplicationSummaryResponse.model_construct(
        id=application.id,
        job_id=application.job_id,
        job=job_summary,
        candidate_name=application.candidate_name,
        candidate_email=application.candidate_email,
        status=application.status,
        file_status=application.file_status,
        applied_at=application.applied_at or datetime.now(timezone.utc),
        resume_extraction=re_summary,
        interview=int_summary,
        resume_score=application.resume_score or 0.0,
        composite_score=application.composite_score or 0.0,
        is_owner=is_owner,
        assigned_hr_id=application.hr_id,
        assigned_hr_name=application.hr.full_name if application.hr else None,
        candidate_photo_path=application.candidate_photo_path,
        photo_url=photo_url
    )

def build_application_detail_response(application: Application, current_user_id: Optional[int] = None) -> ApplicationDetailResponse:
    from app.core.storage import get_signed_url, get_public_url
    
    resume_url = None
    if application.resume_file_path:
        resume_url = get_signed_url(settings.supabase_bucket_resumes, application.resume_file_path)
    
    photo_url = None
    if application.candidate_photo_path:
        photo_url = get_signed_url(settings.supabase_bucket_id_photos, application.candidate_photo_path)
             
    id_card_url = None
    if getattr(application, 'id_card_url', None):
        id_card_url = get_signed_url(settings.supabase_bucket_id_cards, application.id_card_url)

    video_url = None
    if application.interview and application.interview.video_recording_path:
        video_url = get_signed_url(settings.supabase_bucket_videos, application.interview.video_recording_path)

    detail = ApplicationDetailResponse.model_validate(application, from_attributes=True)
    raw_notes = application.hr_notes or ""
    degraded = RIMS_EXTRACTION_DEGRADED_MARKER in raw_notes or _heuristic_extraction_degraded(application)
    
    # Safety: Ensure file_status reflects reality (Step 5)
    inferred_file_status = application.file_status or 'active'
    if application.resume_file_path and not resume_url:
        inferred_file_status = 'missing'
    
    detail.file_status = inferred_file_status
    detail.resume_url = resume_url
    detail.photo_url = photo_url
    detail.id_card_url = id_card_url
    detail.video_url = video_url
    detail.extraction_degraded = degraded

    # Ownership Calculation (Simplified)
    if current_user_id:
        detail.is_owner = (application.hr_id == current_user_id)
    detail.assigned_hr_id = application.hr_id
    detail.assigned_hr_name = application.hr.full_name if application.hr else None

    # 7. Fallback Safety (Scenario 5 Guard)
    if not detail.applied_at:
        fallback_date = application.created_at or datetime.now(timezone.utc)
        detail.applied_at = fallback_date
        logger.error(
            f"INTEGRITY_ALERT: App {application.id} missing applied_at. Using fallback.",
            extra={"app_id": application.id, "fallback_date": str(fallback_date)}
        )
    
    return detail

import asyncio
import time
from pathlib import Path

pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

from app.core.config import get_settings
from app.core.idempotency import is_duplicate_request
from app.core.observability import get_request_id, log_json, safe_hash
settings = get_settings()

from app.core.rate_limiter import limiter
from fastapi import Request

router = APIRouter(prefix="/api/applications", tags=["applications"])

@router.get("/failures", response_model=list[ApplicationResponse])
def get_application_failures(
    db: Session = Depends(get_db),
    current_admin: User = Depends(get_current_admin)
):
    """(Control-Level) Get all applications that have failed processing."""
    from app.services.state_machine import CandidateState
    return db.query(Application).filter(
        or_(
            Application.retry_count > 0,
            Application.status == CandidateState.PERMANENT_FAILURE.value
        )
    ).order_by(Application.last_attempt_at.desc()).all()

@router.get("/ranking/{job_id}")
def get_candidate_ranking(
    job_id: int,
    current_user: User = Depends(get_current_hr),
    db: Session = Depends(get_db)
):
    """Get ranked candidates for a specific job (Point 3)"""
    job = db.query(Job).filter(Job.id == job_id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    # Apply visibility isolation
    if current_user.role.lower() == "hr":
        if job.hr_id != current_user.id:
            raise HTTPException(status_code=403, detail="Forbidden: You do not own this job.")
    
    validate_hr_ownership(job, current_user, resource_name="job")

    from app.services.candidate_service import CandidateService
    service = CandidateService(db)
    ranked = service.get_ranked_candidates(job_id)
    
    result = []
    for idx, app in enumerate(ranked):
        result.append({
            "rank": idx + 1,
            "id": app.id,
            "candidate_name": app.candidate_name,
            "composite_score": app.composite_score,
            "recommendation": app.recommendation,
            "status": app.status
        })
    return result


@router.post("/apply", response_model=ApplicationResponse, status_code=status.HTTP_201_CREATED)
@limiter.limit("5/minute")
async def apply_for_job(
    request: Request,
    job_id: int = Form(...),
    candidate_name: str = Form(...),
    candidate_email: Optional[str] = Form(None),
    candidate_phone: Optional[str] = Form(None),
    resume_file: UploadFile = File(...),
    photo_file: Optional[UploadFile] = File(None),
    background_tasks: BackgroundTasks = BackgroundTasks(),
    db: Session = Depends(get_db)
):
    """Apply for a job with resume (Public endpoint)"""
    # Name validation: just ensure it is not empty to support bulk uploads
    if not candidate_name :
        raise HTTPException(
            status_code=400,
            detail="Valid full name required "#(at least two words, containing letters, hyphens, or apostrophes)."
        )

    request_id = None
    ip_address = None
    try:
        request_id = get_request_id(request)
        if request and request.client:
            ip_address = request.client.host
    except Exception:
        request_id = None
        ip_address = None

    # 1. Email/Phone Presence Check (Point 1)
    if not candidate_email and not candidate_phone:
        raise HTTPException(
            status_code=400,
            detail="At least one valid contact method (email or phone) is required."
        )

    # 2. Email Validation & Normalization (Point 2)
    if candidate_email:
        try:
            from app.core.email_utils import validate_email_strict_enterprise

            candidate_email = validate_email_strict_enterprise(
                candidate_email,
                ip=ip_address,
                request_id=request_id,
                logger=logger,
            )
        except ValueError as e:
            raise HTTPException(status_code=400, detail=str(e))
    else:
        candidate_email = None

    # Non-blocking flag for obviously fake/test domains (H-domain hygiene)
    suspicious_email_domain = None
    if candidate_email:
        try:
            # Minimal, explicit list to avoid over-blocking; can be extended safely later.
            DISPOSABLE_DOMAINS = {
                "fengnu.com",
                "mailinator.com",
                "guerrillamail.com",
                "10minutemail.com",
                "tempmail.com",
            }
            _, domain_part = candidate_email.rsplit("@", 1)
            domain_part = domain_part.lower().strip()
            if domain_part in DISPOSABLE_DOMAINS:
                suspicious_email_domain = domain_part
        except Exception:
            suspicious_email_domain = None

    # 3. Phone Validation & Normalization (Points 3, 4)
    from app.core.phone_utils import compute_phone_hash, normalize_phone_digits

    candidate_phone_raw = candidate_phone if candidate_phone else None
    candidate_phone_normalized = None
    phone_error_reason = None
    if candidate_phone:
        candidate_phone_normalized, phone_error_reason = normalize_phone_digits(candidate_phone)

        if candidate_phone_normalized is None and phone_error_reason is not None:
            try:
                from app.core.observability import log_json

                log_json(
                    logger,
                    "phone_validation_rejected",
                    request_id=request_id,
                    endpoint="/api/applications/apply",
                    user_id=None,
                    status=400,
                    level="warning",
                    extra={"reason": phone_error_reason},
                )
            except Exception:
                pass
            if phone_error_reason in ("letters_present", "invalid_length"):
                raise HTTPException(status_code=400, detail="Phone number must be 10–15 digits")
            if phone_error_reason == "invalid_characters":
                raise HTTPException(status_code=400, detail="Phone number must contain only digits (and separators)")
            raise HTTPException(status_code=400, detail="Phone number must be 10–15 digits")

    # Use correctly normalized digits for both storage and hash (Point 1)
    candidate_phone_hash = compute_phone_hash(candidate_phone_normalized) if candidate_phone_normalized else None

    # Check if job exists and is open
    job = db.query(Job).filter(Job.id == job_id, Job.status == "open").first()
    if not job:
        try:
            from app.core.observability import log_json
            log_json(
                logger,
                "validation_failed",
                request_id=request_id,
                endpoint="/api/applications/apply",
                status=404,
                level="warning",
                extra={"module": "applications", "field": "job_id", "reason": "not_found_or_closed", "input_preview": str(job_id)},
            )
        except Exception:
            pass
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Job not found or not open"
        )

    request_id_header = request.headers.get("X-Request-ID")
    idempotency_key = f"{candidate_email or candidate_phone_normalized}:{job_id}"
    if settings.enable_request_id_idempotency and is_duplicate_request(
        request_id=request_id_header,
        scope="applications.apply",
        key=idempotency_key.lower().strip(),
        ttl_seconds=60,
    ):
        existing_idem = (
            db.query(Application)
            .filter(
                Application.job_id == job_id,
                or_(
                    (Application.candidate_email == candidate_email) if candidate_email else False,
                    (Application.candidate_phone_hash == candidate_phone_hash) if candidate_phone_hash else False
                )
            )
            .first()
        )
        if existing_idem:
            log_json(
                logger,
                "apply_idempotent_replay",
                request_id=request_id,
                endpoint="/api/applications/apply",
                level="info",
                extra={
                    "application_id": existing_idem.id,
                    "job_id": job_id,
                    "email_hash": safe_hash(candidate_email),
                },
            )
            return existing_idem
        raise HTTPException(
            status_code=409,
            detail="Duplicate application request detected. Please wait and retry.",
        )

    # 4. Duplicate Identification (Point 7)
    # Block any user who has already applied to this specific job using THIS email OR THIS phone.
    existing_app = db.query(Application).filter(
        Application.job_id == job_id,
        or_(
            (Application.candidate_email == candidate_email) if candidate_email else False,
            (Application.candidate_phone_hash == candidate_phone_hash) if candidate_phone_hash else False,
        ),
    ).first()

    if existing_app:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="You have already applied for this job.",
        )
    
    # 4. Resume Validation
    MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB
    ALLOWED_RESUME_EXTENSIONS = {".pdf", ".docx", ".doc"}
    
    from app.core.resume_upload_utils import (
        generate_hashed_resume_filename,
        get_resume_extension,
        validate_resume_signature,
    )

    resume_ext = get_resume_extension(resume_file.filename)
    if resume_ext not in ALLOWED_RESUME_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid resume file type. Only .pdf, .docx, and .doc are allowed.",
        )
    content = await resume_file.read()
    if not content:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Resume file is empty.")
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="File too large. Maximum size is 5MB.",
        )
    
    ok, reason = validate_resume_signature(resume_ext, content)
    if not ok:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Invalid resume content.")

    # 5. Photo Validation
    photo_content = None
    if photo_file:
        photo_content = await photo_file.read()
        if not photo_content:
            raise HTTPException(status_code=400, detail="Candidate photo is empty.")
        if len(photo_content) > MAX_FILE_SIZE:
            raise HTTPException(status_code=400, detail="Photo too large. Maximum size is 5MB.")

    # 6. Database Entry (Atomic Transaction)
    # We flush first to get the ID, then upload to storage using that ID as a prefix.
    user_agent = request.headers.get("user-agent")
    filename = generate_hashed_resume_filename(
        candidate_email=candidate_email or f"phone_{candidate_phone_normalized}",
        job_id=job_id,
        resume_ext=resume_ext,
        content=content,
    )

    resume_storage_path = None
    photo_storage_path = None

    warning_notes = None
    if suspicious_email_domain:
        warning_notes = f"Warning: Possibly fake/test email domain detected ({suspicious_email_domain})."

    new_application = Application(
        job_id=job_id,
        hr_id=job.hr_id,
        candidate_name=candidate_name,
        candidate_email=candidate_email,
        candidate_phone_normalized=candidate_phone_normalized,
        candidate_phone_raw=candidate_phone_raw,
        candidate_phone_hash=candidate_phone_hash,
        status="applied",
        hr_notes=warning_notes,
        applied_at=datetime.now(timezone.utc),
        resume_status="pending",
    )

    try:
        db.add(new_application)
        db.flush() # Get ID for storage paths

        # Upload Resume to Supabase
        resume_storage_path = f"{new_application.id}/resume_{filename}"
        returned_resume_path = upload_file(settings.supabase_bucket_resumes, resume_storage_path, content, content_type=resume_file.content_type)
        if not returned_resume_path:
            raise Exception("Resume storage upload failed")
        new_application.resume_file_path = returned_resume_path

        # Upload Photo to Supabase
        if photo_content:
            photo_ext = os.path.splitext(photo_file.filename)[1] if photo_file.filename else ".jpg"
            photo_storage_path = f"{new_application.id}/photo_initial{photo_ext}"
            returned_photo_path = upload_file(settings.supabase_bucket_id_photos, photo_storage_path, photo_content, content_type=photo_file.content_type)
            if not returned_photo_path:
                raise Exception("Candidate photo storage upload failed")
            new_application.candidate_photo_path = returned_photo_path

        # Create HR Notification
        from app.domain.models import Notification
        db.add(Notification(
            user_id=job.hr_id,
            notification_type="new_application",
            title=f"New Application: {candidate_name}",
            message=f"{candidate_name} has applied for {job.title}.",
            related_application_id=new_application.id,
        ))

        db.commit()
        db.refresh(new_application)

        # Audit Log
        from app.services.candidate_service import CandidateService
        CandidateService(db).create_audit_log(None, "APPLICATION_CREATED", "Application", new_application.id, {"job_id": job_id})

        # [Trigger] Application Received Email
        try:
            if not new_application or not new_application.candidate_email:
                logger.error(f"[EMAIL][FAILED] Missing email for App #{getattr(new_application, 'id', 'UNKNOWN')}")
            elif getattr(new_application, "_email_sent", False):
                logger.warning(f"[EMAIL][SKIPPED] Duplicate prevented for App #{new_application.id}")
            else:
                background_tasks.add_task(send_application_received_email, new_application)
                new_application._email_sent = True
                logger.info(f"[EMAIL] Application email queued for App #{new_application.id}")
        except Exception as e:
            logger.error(f"[EMAIL][FAILED] Application email for App #{getattr(new_application, 'id', 'UNKNOWN')}: {str(e)}")


    except Exception as e:
        db.rollback()
        logger.error(f"Application submission failed: {e}")
        # Clean up cloud files if they were uploaded
        from app.core.storage import delete_file
        try:
            if resume_storage_path and new_application.resume_file_path:
                delete_file(settings.supabase_bucket_resumes, resume_storage_path)
            if photo_storage_path and new_application.candidate_photo_path:
                delete_file(settings.supabase_bucket_id_photos, photo_storage_path)
        except Exception as e:
            logger.warning(f"Failed to clean up files for application: {e}")
        raise HTTPException(status_code=500, detail="Failed to submit application securely.")

    background_tasks.add_task(
        process_application_background, 
        new_application.id, 
        job_id, 
        new_application.resume_file_path, 
        candidate_email, 
        candidate_name
    )
    
    return new_application

async def process_application_background(application_id: int, job_id: int, abs_file_path: str, candidate_email: str, candidate_name: str):
    """Heavy AI processing and notification workflow in background"""
    db = SessionLocal()
    try:
        from app.services.candidate_service import CandidateService
        cand_service = CandidateService(db)
        
        # Reload objects in this session
        # Step 1: lock only (without joins)
        # Note: Application.resume_extraction is configured as lazy='joined', so ORM-level
        # .with_for_update() may turn into a LEFT OUTER JOIN, which Postgres rejects.
        db.execute(text("SELECT 1 FROM applications WHERE id = :id FOR UPDATE"), {"id": application_id})
        
        # Step 2: fetch with joins, no lock
        application = db.query(Application).options(
            joinedload(Application.resume_extraction)
        ).filter(Application.id == application_id).first()
        
        job = db.query(Job).filter(Job.id == job_id).first()
        if not application or not job:
            db.close()
            return

        application.resume_status = "parsing"
        application.parsing_started_at = datetime.now(timezone.utc)
        db.commit()
        db.refresh(application)
        try:
            log_json(
                logger,
                "resume_parsing_started",
                level="info",
                extra={"application_id": application_id, "job_id": job_id},
            )
        except Exception:
            pass

        # 1. Initial State
        cand_service.advance_stage(application_id, "Application Submitted", "pass")
        cand_service.create_audit_log(None, "APPLICATION_SUBMITTED", "Application", application_id, {"email": candidate_email})
        
        # 2. Screening Stage
        cand_service.advance_stage(application_id, "Resume Screening", "pending")
        
        # Parse resume text based on file type from Supabase Storage
        resume_text = ""
        try:
            if not abs_file_path:
                raise Exception("No resume file path available (storage likely disabled)")
                
            from app.core.storage import get_supabase_client
            supabase = get_supabase_client()
            if not supabase:
                raise Exception("Supabase client not initialized")
            
            # Download from storage
            # Assuming 'resumes' bucket
            bucket_name = settings.supabase_bucket_resumes
            response = supabase.storage.from_(bucket_name).download(abs_file_path)
            if not response:
                raise Exception(f"Failed to download resume from {abs_file_path}")
            
            from io import BytesIO
            file_stream = BytesIO(response)
            
            file_ext = abs_file_path.lower().split('_')[-1].split('.')[-1] if '.' in abs_file_path else 'pdf'
            
            if file_ext == 'pdf':
                if response.startswith(b'%PDF'):
                    from pypdf import PdfReader
                    reader = PdfReader(file_stream)
                    for page in reader.pages:
                        resume_text += page.extract_text() + "\n"
                else:
                    resume_text = response.decode('utf-8', errors='ignore')
            elif file_ext in ['docx', 'doc']:
                import docx
                doc = docx.Document(file_stream)
                for para in doc.paragraphs:
                    resume_text += para.text + "\n"
            else:
                resume_text = response.decode('utf-8', errors='ignore')
        except Exception as e:
            logger.error(f"Background Text Extraction Skipped or Failed: {e}")
            cand_service.create_audit_log(None, "RESUME_TEXT_EXTRACTION_SKIPPED", "Application", application_id, {"reason": str(e)})
            resume_text = "Parsing skipped: Storage unavailable or file missing."
        
        if not resume_text.strip() or resume_text.startswith("Parsing skipped:"):
            # Skip AI if no text or explicitly skipped due to streaming/parsing failure
            extraction_data = {
                "summary": "AI extraction skipped: " + resume_text,
                "skills": ["Unparsable File"],
                "experience": 0,
                "score": 0.0,
                "match_percentage": 0,
                "extraction_degraded": True,
                "is_resume": False,
                "reasoning": {"ai_justification": "System detected a corrupted, missing, or unparsable file. Standard AI evaluation was skipped."}
            }
            extraction_degraded_flag = True
        else:
            # AI Parsing
            extraction_data = await parse_resume_with_ai(resume_text, job_id, job.description, job.experience_level)
            extraction_degraded_flag = extraction_data.pop("extraction_degraded", False)

        # Store extraction (Versioning + Upsert pattern)
        resume_extraction = db.query(ResumeExtraction).filter(ResumeExtraction.application_id == application_id).first()
        if resume_extraction:
            # ── Phase 3: Versioning (Save old record before overwrite) ──
            try:
                version_count = db.query(ResumeExtractionVersion).filter(ResumeExtractionVersion.application_id == application_id).count()
                old_version = ResumeExtractionVersion(
                    application_id=application_id,
                    version_number=version_count + 1,
                    extracted_text=resume_extraction.extracted_text,
                    extracted_skills=resume_extraction.extracted_skills,
                    resume_score=resume_extraction.resume_score
                )
                db.add(old_version)
            except Exception as e:
                logger.warning(f"Failed to version old resume extraction: {e}")
        else:
            resume_extraction = ResumeExtraction(application_id=application_id)
            db.add(resume_extraction)
            
        resume_extraction.extracted_text = resume_text
        resume_extraction.summary = extraction_data.get("summary", "")
        resume_extraction.extracted_skills = json.dumps(extraction_data.get("skills") or [])
        resume_extraction.years_of_experience = extraction_data.get("experience")
        resume_extraction.education = json.dumps(extraction_data.get("education") or [])
        resume_extraction.previous_roles = json.dumps(extraction_data.get("roles") or [])
        resume_extraction.experience_level = extraction_data.get("experience_level")
        resume_extraction.resume_score = extraction_data.get("score", 0)
        resume_extraction.skill_match_percentage = extraction_data.get("match_percentage", 0)
        resume_extraction.reasoning = {"ai_justification": extraction_data.get("reasoning")}
        
        if extraction_data.get("candidate_name"):
            resume_extraction.candidate_name = extraction_data.get("candidate_name")
        if extraction_data.get("email"):
            resume_extraction.email = extraction_data.get("email")
        if extraction_data.get("phone_number"):
            resume_extraction.phone_number = extraction_data.get("phone_number")
        
        # Update Application summary fields
        application.resume_score = extraction_data.get("score", 0)
        
        # ── Phase 7: Scoring Transparency ──
        application.scoring_metadata = json.dumps({
            "logic_version": "v2.0",
            "weights": {"skills": 0.6, "experience": 0.4},
            "recomputed_at": datetime.now(timezone.utc).isoformat(),
            "extraction_degraded": extraction_degraded_flag
        })
        
        # ── HYBRID IDENTITY EXTRACTION: AI + regex fallback ──
        from app.services.ai_service import extract_email_regex, extract_phone_regex, extract_name_heuristic
        
        extracted_name = extraction_data.get("candidate_name") or extract_name_heuristic(resume_text)
        extracted_email = extraction_data.get("email") or extract_email_regex(resume_text)
        extracted_phone = extraction_data.get("phone_number") or extract_phone_regex(resume_text)
        
        logger.info(f"[IDENTITY SYNC] App #{application_id} | EXTRACTED: name={extracted_name}, email={extracted_email}, phone={extracted_phone}")
        
        is_duplicate = False
        duplicate_app_id = None
        
        # Check for placeholder status BEFORE any updates
        email_is_placeholder = application.candidate_email and "@batch." in application.candidate_email
        name_is_placeholder = not application.candidate_name or len(application.candidate_name.split()) < 2
        
        # 1. Duplicate Detection via Extracted Email
        if extracted_email:
            norm_email = extracted_email.lower().strip()
            existing_match = db.query(Application).filter(
                Application.job_id == job_id,
                Application.candidate_email == norm_email,
                Application.id != application_id
            ).first()
            
            if existing_match:
                is_duplicate = True
                duplicate_app_id = existing_match.id
                logger.info(f"[IDENTITY SYNC] App #{application_id} Conflict: '{norm_email}' already applied to job {job_id}")
                # Internal-only signal: mark HR notes and pipeline stage as possible duplicate.
                application.hr_notes = (application.hr_notes or "") + f"\nNotice: AI detected this as a duplicate of App ID: {duplicate_app_id} (matched on email in resume)."
            else:
                # Safe to update email if it was missing or obviously wrong (placeholder)
                if not application.candidate_email or email_is_placeholder:
                    application.candidate_email = extracted_email
                    logger.info(f"[IDENTITY SYNC] App #{application_id} | Email updated from resume")

        # 2. Name & Phone Sync (Conservative)
        # We only overwrite if the current value is "weak" (one word name or placeholder from batch)
        if extracted_name and (name_is_placeholder or email_is_placeholder):
            application.candidate_name = extracted_name
            logger.info(f"[IDENTITY SYNC] App #{application_id} | Name updated from resume")
            
        if extracted_phone and not application.candidate_phone:
            application.candidate_phone = extracted_phone
            logger.info(f"[IDENTITY SYNC] App #{application_id} | Phone updated from resume")
            
        if extracted_phone and not application.candidate_phone:
            application.candidate_phone = extracted_phone
            logger.info(f"[IDENTITY SYNC] App #{application_id} | Phone updated from resume")

        if extraction_degraded_flag:
            _append_extraction_degraded_marker(application)

        application.resume_status = "parsed"; application.failure_reason = None
        db.commit()
        db.refresh(application)
        try:
            log_json(
                logger,
                "resume_parsing_completed",
                level="info",
                extra={"application_id": application_id, "resume_status": "parsed"},
            )
        except Exception:
            pass

        # ── Pipeline Advancement ──
        # If duplicate, we mark stage as 'fail' so it's not advanced, but we STILL provide the real score.
        # This is *internal-only* and never surfaces as a 409 or candidate-facing error.
        # If not duplicate, we mark as 'pass' to await HR decision.
        stage_status = "fail" if is_duplicate else "pass"
        stage_note = (
            f"Possible duplicate of App #{duplicate_app_id}"
            if is_duplicate
            else "AI analysis complete — awaiting HR decision"
        )
        
        cand_service.advance_stage(
            application_id, 
            "Resume Screening", 
            stage_status, 
            extraction_data.get("score", 0) * 10, 
            stage_note
        )
        
        # ── Phase 6: Critical Audit Logging ──
        cand_service.create_audit_log(
            None, 
            "RESUME_SCREENING_COMPLETED", 
            "Application", 
            application_id, 
            {"score": extraction_data.get("score", 0), "match": extraction_data.get("match_percentage", 0)},
            is_critical=True
        )
        
        db.commit()
    except Exception as e:
        logger.error(
            f"CRITICAL Background Error processing application {application_id}: {e}",
            exc_info=True,
        )
        db.rollback()
        # Log the critical error
        try:
            cand_service = CandidateService(db) # Re-initialize if needed, or pass db
            cand_service.create_audit_log(None, "BACKGROUND_PROCESSING_FAILED", "Application", application_id, {"error": str(e)})
            # Optionally update application status to indicate processing failed
            application = db.query(Application).filter(Application.id == application_id).with_for_update().first()
            if application:
                application.status = "applied"  # Keep in 'applied' — HR can retry or proceed manually
                application.resume_status = "failed"
                application.retry_count = (application.retry_count or 0) + 1
                application.failure_reason = str(e)[:1000]
                application.last_attempt_at = datetime.now(timezone.utc)

                # Escalation (Phase 5 fix)
                if application.retry_count >= 3:
                     from app.services.state_machine import CandidateState
                     application.status = CandidateState.PERMANENT_FAILURE.value
                     application.failure_reason = "[PERMANENT_FAILURE]: " + application.failure_reason

                # Never write raw exception details (may include internal SQL / query text) into HR-visible notes.
                application.hr_notes = (
                    "AI analysis failed. Please click "
                    "Retry Analysis to reprocess."
                )
                _append_extraction_degraded_marker(application)
                db.commit()
                try:
                    log_json(
                        logger,
                        "resume_parsing_failed",
                        level="error",
                        extra={"application_id": application_id, "resume_status": "failed"},
                    )
                except Exception:
                    pass
        except Exception as log_e:
            logger.error(f"Failed to log critical error for application {application_id}: {log_e}")
    finally:
        db.close()
    
    # The return value of a background task is not used by FastAPI.
    # The original snippet returned new_application, but it's not necessary here.
    # Keeping it for consistency with the provided snippet, but it has no effect.
    return application


@router.get("/has-applied", response_model=HasAppliedResponse)
@limiter.limit("20/minute")
def has_applied_for_job(
    request: Request,
    job_id: int,
    candidate_email: str,
    candidate_phone: Optional[str] = None,
    db: Session = Depends(get_db),
):
    """Return whether a (job_id, candidate_email/phone) application already exists.
    
    CRITICAL: Validates against both email and normalized phone hash to strictly
    enforce one application per person per job.
    """
    from app.core.email_utils import validate_email_strict_enterprise
    from app.core.phone_utils import compute_phone_hash, normalize_phone_digits

    # 1. Normalize Email
    if candidate_email:
        try:
            candidate_email = validate_email_strict_enterprise(
                candidate_email,
                ip=None,
                request_id=None,
                logger=logger,
            )
        except ValueError:
            # If email is invalid, it can't match a stored valid email
            candidate_email = None

    # 2. Normalize and Hash Phone
    candidate_phone_hash = None
    if candidate_phone:
        normalized_digits, _ = normalize_phone_digits(candidate_phone)
        if normalized_digits:
            candidate_phone_hash = compute_phone_hash(normalized_digits)

    # 3. Check for duplicates (OR logic)
    if not candidate_email and not candidate_phone_hash:
        return HasAppliedResponse(hasApplied=False)

    existing = (
        db.query(Application.id)
        .filter(
            Application.job_id == job_id,
            or_(
                (Application.candidate_email == candidate_email) if candidate_email else False,
                (Application.candidate_phone_hash == candidate_phone_hash) if candidate_phone_hash else False,
            ),
        )
        .first()
    )
    return HasAppliedResponse(hasApplied=existing is not None)


@router.get("/pending-count")
def get_pending_applications_count(
    current_user: User = Depends(get_current_hr),
    db: Session = Depends(get_db),
):
    """Sidebar / badges: count applications not in a terminal state, scoped to HR's jobs."""
    # Use func.count() directly — avoids loading full ORM objects just to count them.
    q = db.query(func.count(Application.id)).filter(
        ~Application.status.in_(("hired", "rejected")),
        or_(func.trim(Application.file_status).in_(('active', 'missing')), Application.file_status == None)
    )

    # Apply visibility isolation
    if current_user.role.lower() != "super_admin":
        q = q.filter(Application.hr_id == current_user.id)

    count = q.scalar() or 0
    return {"count": count}


@router.get("", response_model=ApplicationListResponse, response_class=ORJSONResponse)
def get_hr_applications(
    job_id: int = None,
    from_date: str = None,
    to_date: str = None,
    status: str = None,
    time_range: str = None,
    search: str = None,
    skip: int = 0,
    limit: int = 20,
    current_user: User = Depends(get_current_hr),
    db: Session = Depends(get_db)
):
    """Get all applications for HR's jobs (HR only)"""
    t_start = time.perf_counter()
    items = []
    total = 0
    safe_skip = max(0, int(skip or 0))
    safe_limit = max(1, min(int(limit or 20), 100))
    
    try:
        # 1. Base Query with optimized loading
        query = db.query(Application).outerjoin(Job).options(
            joinedload(Application.job).load_only(Job.id, Job.title, Job.hr_id, Job.status, Job.job_id),
            joinedload(Application.hr).load_only(User.id, User.full_name),
            joinedload(Application.resume_extraction).load_only(
                ResumeExtraction.id, ResumeExtraction.resume_score,
                ResumeExtraction.skill_match_percentage, ResumeExtraction.experience_level,
                ResumeExtraction.summary, ResumeExtraction.extracted_skills,
            ),
            load_only(
                Application.id, Application.candidate_name, Application.candidate_email,
                Application.status, Application.applied_at, Application.file_status,
                Application.candidate_photo_path, Application.resume_file_path,
                Application.resume_score, Application.composite_score
            ),
            joinedload(Application.interview).load_only(Interview.id, Interview.status, Interview.overall_score),
        )

        # 2. Filters
        if job_id:
            query = query.filter(Application.job_id == job_id)

        if status and status != 'all':
            if status == "applied":
                query = query.filter(Application.status.in_(("applied", "submitted")))
            else:
                query = query.filter(Application.status == status)

        if search and str(search).strip():
            term = f"%{search}%"
            query = query.filter(or_(
                Application.candidate_name.ilike(term),
                Application.candidate_email.ilike(term),
                Job.title.ilike(term)
            ))

        if from_date:
            try:
                sd = datetime.strptime(from_date, "%Y-%m-%d").date()
                query = query.filter(func.date(func.timezone("UTC", Application.applied_at)) >= sd)
            except ValueError:
                logger.warning(f"Invalid from_date format: {from_date}")
        if to_date:
            try:
                ed = datetime.strptime(to_date, "%Y-%m-%d").date()
                query = query.filter(func.date(func.timezone("UTC", Application.applied_at)) <= ed)
            except ValueError:
                logger.warning(f"Invalid to_date format: {to_date}")

        # 3. Security
        # Apply visibility isolation: Anyone not a super_admin is restricted to their own apps
        if current_user.role.lower() != "super_admin":
            query = query.filter(Application.hr_id == current_user.id)
        # Super Admin sees all.

        # 4. Retrieval
        total = query.count()
        applications = query.order_by(Application.applied_at.desc()).offset(safe_skip).limit(safe_limit).all()

        # Path sanitization
        for app in applications:
            for field in ['candidate_photo_path', 'resume_file_path']:
                val = getattr(app, field)
                if val and "uploads" in val:
                    idx = val.find("uploads")
                    setattr(app, field, val[idx:].replace("\\", "/"))

        # 5. Response Mapping
        t_map_start = time.perf_counter()
        items = [build_application_summary_response(app, current_user.id) for app in applications]
        t_map_duration = time.perf_counter() - t_map_start
        
        duration = time.perf_counter() - t_start
        logger.info(f"PERFORMANCE_TRACE: get_hr_applications total={duration:.4f}s (Map: {t_map_duration:.4f}s)")
        
        pages = (total + safe_limit - 1) // safe_limit
        return {
            "items": items,
            "total": total,
            "page": (safe_skip // safe_limit) + 1,
            "size": safe_limit,
            "pages": pages
        }
    except Exception as e:
        logger.error(f"APPLICATION_API_ERROR: {str(e)}", exc_info=True)
        return {"items": [], "total": 0, "page": 1, "size": safe_limit, "pages": 0, "error_hint": str(e)}

@router.get("/{application_id}/resume/download")
def download_resume(
    application_id: int,
    request: Request,
    current_user: User = Depends(get_current_hr),
    db: Session = Depends(get_db)
):
    """Securely download a candidate's resume (HR only)"""
    application = db.query(Application).filter(Application.id == application_id).with_for_update().first()
    if not application or not application.resume_file_path:
        raise HTTPException(status_code=404, detail="Resume file not found")
    validate_hr_ownership(application, current_user, resource_name="application")

    from fastapi.responses import RedirectResponse, FileResponse
    stored_path = (application.resume_file_path or "").replace("\\", "/")
    
    # 1. Cloud Storage Redirect (New)
    if "uploads" not in stored_path:
        from app.core.storage import get_signed_url
        signed_url = get_signed_url(settings.supabase_bucket_resumes, application.resume_file_path)
        if signed_url:
            return RedirectResponse(url=signed_url)
    
    # 2. Local File Fallback (Legacy)
    filename = os.path.basename(stored_path)
    candidate_1 = settings.uploads_dir / "resumes" / filename
    candidate_2 = None
    if "uploads/" in stored_path:
        rel = stored_path.split("uploads/", 1)[1]
        candidate_2 = settings.uploads_dir / rel
    candidate_3 = settings.uploads_dir / filename

    file_path = None
    for c in [candidate_1, candidate_2, candidate_3]:
        if c and c.exists():
            file_path = c
            break

    if not file_path:
        raise HTTPException(status_code=404, detail="Resume file not found on server")

    return FileResponse(
        path=str(file_path),
        filename=application.resume_file_name or filename,
        media_type='application/octet-stream'
    )

@router.get("/{application_id}", response_model=ApplicationDetailResponse)
def get_application(
    application_id: int,
    current_user: User = Depends(get_current_hr),
    db: Session = Depends(get_db)
):
    """Get application details (HR only)"""
    application = db.query(Application).options(
        joinedload(Application.job),
        joinedload(Application.hr),
        joinedload(Application.resume_extraction),
        joinedload(Application.interview),
        selectinload(Application.pipeline_stages)
    ).filter(Application.id == application_id).first()
    
    if not application:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Application not found"
        )
    validate_hr_ownership(application, current_user, resource_name="application")
    
    # Sanitize paths
    if application.candidate_photo_path and ":" in application.candidate_photo_path:
        idx = application.candidate_photo_path.find("uploads")
        if idx != -1:
            application.candidate_photo_path = application.candidate_photo_path[idx:].replace("\\", "/")
    if application.resume_file_path and ":" in application.resume_file_path:
        idx = application.resume_file_path.find("uploads")
        if idx != -1:
            application.resume_file_path = application.resume_file_path[idx:].replace("\\", "/")

    return build_application_detail_response(application, current_user.id)

async def retry_application_background(application_id: int, job_id: int, bucket_path: str):
    """Safely retry AI resume extraction without altering pipeline stages or triggering emails."""
    db = SessionLocal()
    try:
        from app.services.candidate_service import CandidateService
        cand_service = CandidateService(db)
        
        # Reload objects in this session
        # Step 1: lock only (without joins)
        # Note: Application.resume_extraction is configured as lazy='joined', so ORM-level
        # .with_for_update() may turn into a LEFT OUTER JOIN, which Postgres rejects.
        db.execute(text("SELECT 1 FROM applications WHERE id = :id FOR UPDATE"), {"id": application_id})
        
        # Step 2: fetch with joins, no lock
        application = db.query(Application).options(
            joinedload(Application.resume_extraction)
        ).filter(Application.id == application_id).first()
        
        job = db.query(Job).filter(Job.id == job_id).first()
        if not application or not job:
            db.close()
            return
            
        logger.info(f"Retrying AI extraction for application {application_id}...")
        
        # Parse resume text based on file type from Supabase
        resume_text = ""
        try:
            from io import BytesIO
            from app.core.storage import get_supabase_client
            supabase = get_supabase_client()
            response = supabase.storage.from_(settings.supabase_bucket_resumes).download(bucket_path)
            file_stream = BytesIO(response)
            file_ext = bucket_path.lower().split('_')[-1].split('.')[-1] if '.' in bucket_path else 'pdf'

            if file_ext == 'pdf':
                from pypdf import PdfReader
                reader = PdfReader(file_stream)
                for page in reader.pages:
                    resume_text += page.extract_text() + "\n"
            elif file_ext in ['docx', 'doc']:
                import docx
                doc = docx.Document(file_stream)
                for para in doc.paragraphs:
                    resume_text += para.text + "\n"
            else:
                file_stream.seek(0)
                resume_text = file_stream.read().decode('utf-8', errors='ignore')
        except Exception as e:
            logger.error(f"Retry Text Extraction Error: {e}", exc_info=True)
            cand_service.create_audit_log(None, "RETRY_TEXT_EXTRACTION_FAILED", "Application", application_id, {"error": str(e)})
            resume_text = "Error extracting text."
        
        if not resume_text.strip():
            resume_text = "No readable text found."

        # AI Parsing
        extraction_data = await parse_resume_with_ai(resume_text, job_id, job.description, job.experience_level)
        extraction_degraded_flag = extraction_data.pop("extraction_degraded", False)

        # Look for existing extraction or create new
        resume_extraction = db.query(ResumeExtraction).filter(ResumeExtraction.application_id == application_id).first()
        if not resume_extraction:
            resume_extraction = ResumeExtraction(application_id=application_id)
            db.add(resume_extraction)
            
        resume_extraction.extracted_text = resume_text
        resume_extraction.summary = extraction_data.get("summary", "")
        resume_extraction.extracted_skills = json.dumps(extraction_data.get("skills") or [])
        resume_extraction.years_of_experience = extraction_data.get("experience")
        resume_extraction.education = json.dumps(extraction_data.get("education") or [])
        resume_extraction.previous_roles = json.dumps(extraction_data.get("roles") or [])
        resume_extraction.experience_level = extraction_data.get("experience_level")
        resume_extraction.resume_score = extraction_data.get("score", 0)
        resume_extraction.skill_match_percentage = extraction_data.get("match_percentage", 0)
        
        if extraction_data.get("candidate_name"):
            resume_extraction.candidate_name = extraction_data.get("candidate_name")
        if extraction_data.get("email"):
            resume_extraction.email = extraction_data.get("email")
        if extraction_data.get("phone_number"):
            resume_extraction.phone_number = extraction_data.get("phone_number")
        
        application.resume_score = extraction_data.get("score", 0)
        
        if "@batch." in application.candidate_email:
            if extraction_data.get("candidate_name"):
                application.candidate_name = extraction_data.get("candidate_name")
            if extraction_data.get("email"):
                application.candidate_email = extraction_data.get("email")
            if extraction_data.get("phone_number"):
                application.candidate_phone = extraction_data.get("phone_number")

        if extraction_degraded_flag:
            _append_extraction_degraded_marker(application)

        application.resume_status = "parsed"; application.failure_reason = None
        db.commit()
        cand_service.create_audit_log(None, "AI_ANALYSIS_RETRY_SUCCESS", "Application", application_id, {"score": extraction_data.get("score")})
        logger.info(f"Retry successful for application {application_id}")
        try:
            log_json(
                logger,
                "resume_retry_background_completed",
                level="info",
                extra={"application_id": application_id, "resume_status": "parsed"},
            )
        except Exception:
            pass
    except Exception as e:
        logger.error(
            f"Retry Background Error processing application {application_id}: {e}",
            exc_info=True,
        )
        db.rollback()
        try:
            cand_service = CandidateService(db)
            cand_service.create_audit_log(None, "AI_ANALYSIS_RETRY_FAILED", "Application", application_id, {"error": str(e)})
            failed_app = db.query(Application).filter(Application.id == application_id).first()
            if failed_app:
                failed_app.resume_status = "failed"
                failed_app.retry_count = (failed_app.retry_count or 0) + 1
                failed_app.failure_reason = str(e)[:1000] # Cap length
                # Never write raw exception details (may include internal SQL / query text) into HR-visible notes.
                failed_app.hr_notes = (
                    "AI analysis failed. Please click "
                    "Retry Analysis to reprocess."
                )
                db.commit()
            try:
                log_json(
                    logger,
                    "resume_retry_background_failed",
                    level="error",
                    extra={"application_id": application_id, "resume_status": "failed"},
                )
            except Exception:
                pass
        except Exception:
            pass
    finally:
        db.close()

@router.post("/{application_id}/retry-analysis")
async def retry_resume_analysis(
    application_id: int, 
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_hr),
    db: Session = Depends(get_db)
):
    """Manually trigger AI resume analysis if it failed"""
    application = db.query(Application).filter(Application.id == application_id).with_for_update().first()
    if not application:
        raise HTTPException(status_code=404, detail="Application not found")
    validate_hr_ownership(application, current_user, resource_name="application")
        
    application.resume_status = "parsing"
    application.updated_at = datetime.now(timezone.utc)
    db.commit()
    db.refresh(application)
    try:
        log_json(
            logger,
            "resume_retry_api_accepted",
            level="info",
            extra={
                "application_id": application_id,
                "hr_user_id": current_user.id,
                "resume_status": "parsing",
            },
        )
    except Exception:
        pass

    background_tasks.add_task(
        retry_application_background, 
        application.id, 
        application.job_id, 
        application.resume_file_path
    )
    
    return {
        "status": "success",
        "message": "Analysis restarted in background safely",
        "application_id": application_id,
        "resume_status": "parsing",
    }

@router.post("/{application_id}/resend-interview-invitation")
async def resend_interview_invitation(
    application_id: int,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_hr),
    db: Session = Depends(get_db),
):
    """
    Re-send the interview invitation email.

    This is useful when the original "approve_for_interview" succeeded but the email
    provider (e.g. Gmail quota) failed. We re-issue the interview access key only when
    the interview is missing or still 'not_started'.
    """
    # Concurrency Hardening (Phase 1): Lock row for update
    application = db.query(Application).options(
        joinedload(Application.job),
        joinedload(Application.interview),
    ).filter(Application.id == application_id).with_for_update().first()

    if not application:
        raise HTTPException(status_code=404, detail="Application not found")

    validate_hr_ownership(application, current_user, resource_name="application")

    # If an interview exists and is already in progress/completed, don't re-issue the key.
    if application.interview and getattr(application.interview, "status", None) != "not_started":
        raise HTTPException(
            status_code=400,
            detail="Interview access key cannot be reissued unless interview is 'not_started'",
        )

    from app.services.candidate_service import CandidateService
    cand_service = CandidateService(db)
    raw_access_key = cand_service.ensure_interview_record_exists(application)
    candidate_email = application.candidate_email
    job_title = application.job.title if application.job else "your applied position"

    background_tasks.add_task(
        send_approved_for_interview_email,
        candidate_email,
        job_title,
        raw_access_key,
    )

    try:
        log_json(
            logger,
            "resume_invite_email_resend_scheduled",
            level="info",
            extra={"application_id": application_id, "to": candidate_email},
        )
    except Exception:
        pass

    return {
        "status": "success",
        "message": "Interview invitation email scheduled in background",
        "application_id": application_id,
    }

@router.put("/{application_id}/status")
async def update_application_status(
    application_id: int,
    status_update: ApplicationStatusUpdate,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_hr),
    db: Session = Depends(get_db)
):
    """
    Execute a state transition on an application via the finite state machine.
    
    The frontend sends an 'action' (e.g. 'approve_for_interview', 'reject', 
    'call_for_interview', 'review_later', 'hire').  The FSM validates the 
    transition, updates the status atomically, logs the change, and returns 
    the result including which email to send.
    """
    from app.services.state_machine import (
        CandidateStateMachine, TransitionAction,
        InvalidTransitionError, DuplicateTransitionError,
    )

    application = db.query(Application).filter(Application.id == application_id).with_for_update().first()
    if not application:
        raise HTTPException(status_code=404, detail="Application not found")
    validate_hr_ownership(application, current_user, resource_name="application")

    # Parse the action
    try:
        action = TransitionAction(status_update.action)
    except ValueError:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid action: '{status_update.action}'. "
                   f"Valid actions: {[a.value for a in TransitionAction if not a.value.startswith('system_')]}"
        )

    # Block system actions from HR endpoint
    if action.value.startswith("system_"):
        raise HTTPException(status_code=400, detail="System actions cannot be triggered manually")

    # Execute FSM transition
    fsm = CandidateStateMachine(db)
    try:
        logger.debug(f"/api/auth/me user_id={current_user.id}, role={current_user.role}")
        # Hardening Phase 5: Single transaction for Status + Audit Log
        result = fsm.transition(
            application=application,
            action=action,
            user_id=current_user.id,
            notes=status_update.hr_notes,
            is_critical=True,
            background_tasks=background_tasks
        )
        # Flush to confirm DB state without committing yet
        db.flush()
    except InvalidTransitionError as e:
        raise HTTPException(status_code=400, detail=e.message)
    except DuplicateTransitionError as e:
        raise HTTPException(status_code=400, detail=str(e))

    # Handle HR notes
    if status_update.hr_notes:
        application.hr_notes = status_update.hr_notes

    # ─── Post-transition side effects ───────────────────────────────────
    raw_access_key = None

    if action == TransitionAction.APPROVE_FOR_INTERVIEW:
        # Create or refresh interview record + access key
        from app.services.candidate_service import CandidateService
        cand_service = CandidateService(db)
        raw_access_key = cand_service.ensure_interview_record_exists(application)

    if action == TransitionAction.HIRE:
        # Create HiringDecision record
        from app.domain.models import HiringDecision
        existing = db.query(HiringDecision).filter(
            HiringDecision.application_id == application_id
        ).first()
        if not existing:
            decision = HiringDecision(
                application_id=application_id,
                hr_id=current_user.id,
                decision="hired",
                decision_comments=status_update.hr_notes or "Hired via pipeline",
                decided_at=datetime.now(timezone.utc),
            )
            db.add(decision)

    if action == TransitionAction.REJECT:
        from app.domain.models import HiringDecision
        existing = db.query(HiringDecision).filter(
            HiringDecision.application_id == application_id
        ).first()
        if not existing:
            decision = HiringDecision(
                application_id=application_id,
                hr_id=current_user.id,
                decision="rejected",
                decision_comments=status_update.hr_notes or "Rejected via pipeline",
                decided_at=datetime.now(timezone.utc),
            )
            db.add(decision)
    # ────────────────────────────────────────────────────────────────────

    # Atomic commit
    try:
        db.commit()
        db.refresh(application)
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail="Failed to save state transition")

    # ─── Email triggers (ONLY after successful commit) ──────────────────
    candidate_email = application.candidate_email
    job_title = application.job.title

    logger.info(
        f"EMAIL_TRIGGER_CHECK: app={application.id}, "
        f"email_type={result.email_type}, "
        f"raw_access_key_exists={raw_access_key is not None}, "
        f"candidate_email={candidate_email}"
    )

    if result.email_type == "approved_for_interview" and raw_access_key:
        logger.info(f"[EMAIL] Scheduling approved_for_interview email to {candidate_email}")
        background_tasks.add_task(send_approved_for_interview_email, candidate_email, job_title, raw_access_key)
    elif result.email_type == "rejected":
        logger.info(f"[EMAIL] Scheduling rejected email to {candidate_email}")
        background_tasks.add_task(send_rejected_email, candidate_email, job_title, False)
    elif result.email_type == "call_for_interview":
        logger.info(f"[EMAIL] Scheduling call_for_interview email to {candidate_email}")
        from app.services.email_service import send_call_for_interview_email
        background_tasks.add_task(send_call_for_interview_email, candidate_email, job_title)
    elif result.email_type == "hired":
        logger.info(f"[EMAIL] Scheduling hired email to {candidate_email}")
        from app.services.email_service import send_hired_email
        background_tasks.add_task(send_hired_email, candidate_email, job_title, application.interview)
    elif result.email_type:
        logger.warning(f"[EMAIL] No email trigger matched for email_type={result.email_type}")
    # ────────────────────────────────────────────────────────────────────

    return {
        "id": application.id,
        "status": application.status,
        "transition": {
            "from_state": result.from_state,
            "to_state": result.to_state,
            "action": result.action,
            "email_type": result.email_type,
        }
    }



@router.delete("/{application_id}")
async def delete_application(
    application_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_hr)
):
    """
    Delete an application along with associated data. HR only.
    """
    app = db.query(Application).filter(Application.id == application_id).first()
    if not app:
        raise HTTPException(status_code=404, detail="Application not found")
    validate_hr_ownership(app, current_user, resource_name="application")

    try:
        # Explicitly delete ResumeExtraction to prevent ForeignKeyViolation
        # (Table missing ON DELETE CASCADE and relationship has passive_deletes=True)
        from app.domain.models import ResumeExtraction
        db.query(ResumeExtraction).filter(ResumeExtraction.application_id == application_id).delete()
        
        db.delete(app)
        db.commit()
    except Exception as e:
        db.rollback()
        logger.error(f"Error deleting application {application_id}: {e}")
        raise HTTPException(status_code=500, detail="Failed to delete application. It might have complex dependencies.")
    return {"message": "Application deleted successfully"}

@router.post("/{application_id}/merge/{target_id}")
async def merge_applications(
    application_id: int,
    target_id: int,
    current_user: User = Depends(get_current_hr),
    db: Session = Depends(get_db)
):
    """
    Merge a source application into a target application.
    Used for resolving duplicate submissions.
    """
    source = db.query(Application).filter(Application.id == application_id).first()
    target = db.query(Application).filter(Application.id == target_id).first()
    
    if not source or not target:
        raise HTTPException(status_code=404, detail="One or both applications not found")
    validate_hr_ownership(source, current_user, resource_name="application")
    validate_hr_ownership(target, current_user, resource_name="application")
        
    if source.job_id != target.job_id:
        raise HTTPException(status_code=400, detail="Applications must belong to the same job to be merged")
        
    # Merge strategy: Target keeps its identity, but takes scores/notes from source if they are better
    try:
        if source.resume_score > (target.resume_score or 0):
            target.resume_score = source.resume_score
            
        target.hr_notes = (target.hr_notes or "") + f"\n[MERGED from App #{application_id}]: " + (source.hr_notes or "No notes")
        
        # Log the merge
        from app.domain.models import AuditLog
        merge_log = AuditLog(
            user_id=current_user.id,
            action="APPLICATION_MERGED",
            resource_type="Application",
            resource_id=target.id,
            details=json.dumps({"source_id": application_id, "target_id": target_id})
        )
        db.add(merge_log)
        
        # Mark source as rejected/duplicate and hide it or delete it
        source.status = "rejected"
        source.hr_notes = (source.hr_notes or "") + f"\n[MERGED INTO App #{target_id}]"
        
        db.commit()
    except Exception as e:
        db.rollback()
        logger.error(f"Merge error: {e}")
        raise HTTPException(status_code=500, detail="Failed to merge applications")
        
    return {"status": "success", "message": f"Application {application_id} merged into {target_id}"}
@router.put("/{application_id}/notes", response_model=ApplicationResponse)
async def update_hr_notes(
    application_id: int,
    notes_update: ApplicationNotesUpdate,
    current_user: User = Depends(get_current_hr),
    db: Session = Depends(get_db)
):
    """Update HR notes for an application"""
    application = db.query(Application).filter(Application.id == application_id).with_for_update().first()
    if not application:
        raise HTTPException(status_code=404, detail="Application not found")
    validate_hr_ownership(application, current_user, resource_name="application")
        
    application.hr_notes = notes_update.hr_notes
    db.commit()
    db.refresh(application)
    return application

@router.post("/extract-basic-info")
async def extract_basic_info(resume_file: UploadFile = File(...)):
    """Fast endpoint to extract Name and Phone from an uploaded resume."""
    # Read file content
    content = await resume_file.read()
    
    # Extract text locally in memory
    resume_text = ""
    file_ext = resume_file.filename.lower().split('.')[-1]
    
    try:
        if file_ext == 'pdf':
            from pypdf import PdfReader
            import io
            reader = PdfReader(io.BytesIO(content))
            for page in reader.pages:
                resume_text += page.extract_text() + "\n"
        elif file_ext in ['docx', 'doc']:
            import docx
            import io
            doc = docx.Document(io.BytesIO(content))
            for para in doc.paragraphs:
                resume_text += para.text + "\n"
        else:
            resume_text = content.decode('utf-8', errors='ignore')
    except Exception as e:
        logger.error(f"Error extracting text for basic info: {e}")
        return {"name": "", "phone": ""}
        
    if not resume_text.strip():
        return {"name": "", "phone": ""}
        
    info = await extract_basic_candidate_info(resume_text)
    # Privacy & correctness: never return or pre-fill email here.
    # Only expose minimal fields required for UX pre-fill.
    return {
        "name": info.get("name") if isinstance(info, dict) else "",
        "phone": info.get("phone") if isinstance(info, dict) else "",
    }
